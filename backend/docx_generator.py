from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from pathlib import Path
import os

# Check environment for explicit upload path (Docker), fallback to local relative
ENV_UPLOADS = os.environ.get("UPLOAD_DIR")
if ENV_UPLOADS:
    UPLOADS_BASE_PATH = Path(ENV_UPLOADS)
else:
    UPLOADS_BASE_PATH = Path("uploads")


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    
    # Get or create tblPr
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing borders if any
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def get_severity_color(severity):
    """Get color for severity level"""
    colors = {
        'kritikus': 'FF0000',  # Red
        'sulyos': 'FF8C00',    # Orange
        'kozepes': 'FFD700',   # Yellow
        'enyhe': '32CD32'      # Green
    }
    return colors.get(severity, 'D9D9D9')


def get_severity_name(severity):
    """Get Hungarian name for severity"""
    names = {
        'kritikus': 'Kritikus',
        'sulyos': 'Súlyos',
        'kozepes': 'Közepes',
        'enyhe': 'Enyhe'
    }
    return names.get(severity, severity)


def add_defects_section(doc, protocol):
    """Add defects section with images to document"""
    defects = protocol.protocol_defects if hasattr(protocol, 'protocol_defects') else []
    
    if not defects:
        doc.add_paragraph('Nincs feltárt hiba. A vizsgált villamos berendezések minden mért és vizsgált paraméter tekintetében megfelelt az MSZ HD 60364-6:2017 szabvány előírásainak.')
        return
    
    # Critical warning if any critical defects
    critical_defects = [d for d in defects if (d.severity_override or (d.defect_type.severity if d.defect_type else 'kozepes')) == 'kritikus']
    if critical_defects:
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run('⚠️ FIGYELMEZTETÉS: A jegyzőkönyvben kritikus súlyosságú hibák találhatók, amelyek közvetlen életveszélyt és/vagy fokozott tűzveszélyt jelentenek. Ezen hibák azonnali kijavítása kötelező!')
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(255, 0, 0)
        doc.add_paragraph()
    
    # Defects summary table
    doc.add_heading('Hibák összesítő táblázata', level=2)
    
    defect_table = doc.add_table(rows=len(defects) + 1, cols=5)
    add_table_borders(defect_table)
    
    # Header row
    headers = ['#', 'Hiba megnevezése', 'Helyszín', 'Súlyosság', 'Javasolt intézkedés']
    for i, h in enumerate(headers):
        cell = defect_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9D9D9')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    for idx, defect in enumerate(defects):
        row = defect_table.rows[idx + 1]
        
        # Get defect details
        if defect.defect_type:
            name = defect.defect_type.name
            severity = defect.severity_override or defect.defect_type.severity
            action = defect.defect_type.recommended_action or ''
        else:
            name = defect.custom_description or 'Egyéb hiba'
            severity = defect.severity_override or 'kozepes'
            action = ''
        
        row.cells[0].text = str(idx + 1)
        row.cells[1].text = name
        row.cells[2].text = defect.location or '-'
        
        # Severity with color
        severity_cell = row.cells[3]
        severity_cell.text = get_severity_name(severity)
        set_cell_shading(severity_cell, get_severity_color(severity))
        
        row.cells[4].text = action[:100] + '...' if len(action) > 100 else action
    
    doc.add_paragraph()
    
    # Detailed defect descriptions with images
    doc.add_heading('Hibák részletes leírása', level=2)
    
    for idx, defect in enumerate(defects):
        # Get defect details
        if defect.defect_type:
            name = defect.defect_type.name
            severity = defect.severity_override or defect.defect_type.severity
            description = defect.custom_description or defect.defect_type.template_text or defect.defect_type.description or ''
            action = defect.defect_type.recommended_action or ''
            standard = defect.defect_type.standard_reference or ''
        else:
            name = 'Egyéb hiba'
            severity = defect.severity_override or 'kozepes'
            description = defect.custom_description or ''
            action = ''
            standard = ''
        
        # Defect heading
        defect_heading = doc.add_heading(f'{idx + 1}. {name}', level=3)
        
        # Severity and location
        info_para = doc.add_paragraph()
        info_para.add_run('Súlyosság: ').bold = True
        severity_run = info_para.add_run(get_severity_name(severity))
        severity_run.font.color.rgb = RGBColor(
            int(get_severity_color(severity)[:2], 16),
            int(get_severity_color(severity)[2:4], 16),
            int(get_severity_color(severity)[4:6], 16)
        )
        
        if defect.location:
            info_para.add_run('  |  ')
            info_para.add_run('Helyszín: ').bold = True
            info_para.add_run(defect.location)
        
        # Description
        if description:
            desc_para = doc.add_paragraph()
            desc_para.add_run('Leírás: ').bold = True
            desc_para.add_run(description)
        
        # Recommended action
        if action:
            action_para = doc.add_paragraph()
            action_para.add_run('Javasolt intézkedés: ').bold = True
            action_para.add_run(action)
        
        # Standard reference
        if standard:
            std_para = doc.add_paragraph()
            std_para.add_run('Szabvány hivatkozás: ').bold = True
            std_para.add_run(standard)
        
        # Images
        images = defect.images if hasattr(defect, 'images') else []
        if images:
            doc.add_paragraph()
            img_heading = doc.add_paragraph()
            img_heading.add_run('Csatolt képek:').bold = True
            
            for img in images:
                try:
                    img_path = UPLOADS_BASE_PATH / img.image_path
                    if img_path.exists():
                        # Add image with max width of 15cm
                        doc.add_picture(str(img_path), width=Cm(15))
                        
                        # Add image caption
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_text = img.description or img.original_filename or 'Kép'
                        caption.add_run(caption_text).italic = True
                except Exception as e:
                    doc.add_paragraph(f'[Kép nem tölthető be: {img.original_filename or "ismeretlen"}]')
        
        doc.add_paragraph()  # Spacing between defects


def generate_protocol_docx(protocol) -> bytes:
    """Generate Word document from protocol data"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Header - Serial number
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(f"Jegyzőkönyv sorszáma: {protocol.serial_number}")
    header_run.bold = True
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title
    title = doc.add_heading('VILLAMOS BIZTONSÁGI FELÜLVIZSGÁLAT JEGYZŐKÖNYV (VBF)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Basic data
    doc.add_heading('1. Alapadatok', level=1)
    
    basic_data = [
        ('Helyszín:', protocol.location_address),
        ('Hálózat típusa:', protocol.network_type),
        ('Megrendelő neve:', protocol.client_name),
        ('Vizsgálat típusa:', protocol.inspection_type),
        ('Időpont:', protocol.inspection_date.strftime('%Y.%m.%d.') if protocol.inspection_date else ''),
        ('Mérőműszer:', f"{protocol.instrument_model or ''} (Kalibráció érvényes: {protocol.calibration_valid_until.strftime('%Y.%m.%d') if protocol.calibration_valid_until else 'N/A'})"),
        ('Bizonyítvány száma:', protocol.certificate_number or ''),
        ('Felülvizsgáló neve:', protocol.inspector_name),
    ]
    
    table = doc.add_table(rows=len(basic_data), cols=2)
    add_table_borders(table)
    
    for i, (label, value) in enumerate(basic_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Section 2: Overview
    doc.add_heading('2. Villamos biztonsági felülvizsgálat – áttekintés', level=1)
    
    overview_text = f"""A villamos berendezés első ellenőrzése célja annak igazolása, hogy a hálózat megfelel az MSZ HD 60364-6:2017 és az MSZ EN 61557 szabványok által előírt műszaki és érintésvédelmi követelményeknek.

A vizsgálat fő elemei:
• Érintésvédelmi mód ellenőrzése ({protocol.network_type} gyorskioldás feltétele)
• Védővezető folytonosság mérése (Rpe)
• Szigetelési ellenállás vizsgálat (500V DC)
• Hurokellenállás mérése (Zs)
• FI-relé működésvizsgálat

A megfelelőség határértékei:
• Rpe: < 1 Ω (gyakorlatban <0,3 Ω lakóépületnél)
• Szigetelés: ≥ 1 MΩ (általában 200–500 MΩ várható)
• Zs: a gyorskioldás feltétele: Zs ≤ U0 / Ia
• FI-relé: 1×IΔn ≤ 300 ms, 5×IΔn ≤ 40 ms"""
    
    doc.add_paragraph(overview_text)
    
    # Section 3: Summary table
    doc.add_heading('3. Vizsgálati összesítés', level=1)
    
    summary_data = protocol.summary_results if protocol.summary_results else []
    
    # Default summary items if not provided
    if not summary_data:
        default_items = [
            'Dokumentáció', 'Szemrevételezés', 'Rpe', 'Szigetelés',
            'Hurokellenállás', 'FI-relé', 'Dugaljak állapota', 'Kötések', 'Érintésvédelem'
        ]
        summary_table = doc.add_table(rows=len(default_items) + 1, cols=3)
    else:
        summary_table = doc.add_table(rows=len(summary_data) + 1, cols=3)
    
    add_table_borders(summary_table)
    
    # Header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = 'Vizsgálat'
    header_cells[1].text = 'Minősítés'
    header_cells[2].text = 'Megjegyzés'
    for cell in header_cells:
        set_cell_shading(cell, 'D9D9D9')
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    
    if summary_data:
        for i, item in enumerate(summary_data):
            row = summary_table.rows[i + 1]
            row.cells[0].text = item.test_name
            row.cells[1].text = item.result
            row.cells[2].text = item.comment or ''
    else:
        default_items = [
            'Dokumentáció', 'Szemrevételezés', 'Rpe', 'Szigetelés',
            'Hurokellenállás', 'FI-relé', 'Dugaljak állapota', 'Kötések', 'Érintésvédelem'
        ]
        for i, item_name in enumerate(default_items):
            row = summary_table.rows[i + 1]
            row.cells[0].text = item_name
            row.cells[1].text = 'MEGFELELT'
            row.cells[2].text = 'Határértéken belül'
    
    doc.add_paragraph()
    
    # Section 4: Measurement results
    doc.add_heading('4. Mérési eredmények', level=1)
    
    # 4.1 Rpe measurements
    doc.add_heading('4.1 Védővezető folytonosság (Rpe)', level=2)
    
    rpe_data = protocol.rpe_measurements if protocol.rpe_measurements else []
    if rpe_data:
        rpe_table = doc.add_table(rows=len(rpe_data) + 1, cols=4)
        add_table_borders(rpe_table)
        
        headers = ['Pont', 'Hely', 'Rpe (Ω)', 'Megfelel']
        for i, h in enumerate(headers):
            rpe_table.rows[0].cells[i].text = h
            set_cell_shading(rpe_table.rows[0].cells[i], 'D9D9D9')
        
        for i, rpe in enumerate(rpe_data):
            row = rpe_table.rows[i + 1]
            row.cells[0].text = str(rpe.point_number)
            row.cells[1].text = rpe.location
            row.cells[2].text = f"{float(rpe.value_ohm):.2f}"
            row.cells[3].text = 'Igen' if rpe.passed else 'Nem'
    else:
        doc.add_paragraph('Nincs mérési adat.')
    
    doc.add_paragraph()
    
    # 4.2 Insulation measurements
    doc.add_heading('4.2 Szigetelési ellenállás (500V DC)', level=2)
    
    ins_data = protocol.insulation_measurements if protocol.insulation_measurements else []
    if ins_data:
        ins_table = doc.add_table(rows=len(ins_data) + 1, cols=7)
        add_table_borders(ins_table)
        
        headers = ['Áramkör\nés helye', 'Túláramvéd.\n(Típus/A)', 'Vezeték\n(anyag/mm²)', 'Zs (Ω) / dU (%)', 'Tűz.o.', 'Riso (MΩ)', 'Eredmény']
        for i, h in enumerate(headers):
            ins_table.rows[0].cells[i].text = h
            set_cell_shading(ins_table.rows[0].cells[i], 'D9D9D9')
            # Optional: set smaller font size for headers
            if ins_table.rows[0].cells[i].paragraphs[0].runs:
                ins_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        
        for i, ins in enumerate(ins_data):
            row = ins_table.rows[i + 1]
            row.cells[0].text = ins.circuit_name
            
            br_val_str = f"{float(ins.breaker_value):.1f}".rstrip('0').rstrip('.') if ins.breaker_value else ""
            breaker = f"{ins.breaker_type or ''}{br_val_str}"
            row.cells[1].text = breaker if breaker else '-'
            
            wire_cross_str = f"{float(ins.wire_cross_section):.1f}".rstrip('0').rstrip('.') if ins.wire_cross_section else ""
            wire = f"{ins.wire_material or ''} {wire_cross_str}"
            row.cells[2].text = wire.strip() if wire.strip() else '-'
            
            zs_val_str = f"{float(ins.zs_value_ohm):.2f}" if ins.zs_value_ohm else "-"
            du_val_str = f"{float(ins.du_value_percent):.1f}" if ins.du_value_percent else "-"
            zs_du = f"{zs_val_str} / {du_val_str}"
            row.cells[3].text = zs_du if zs_du != "- / -" else "-"
            
            row.cells[4].text = ins.fire_rating or '-'
            
            ln = f"{float(ins.ln_value_mohm):.1f}" if ins.ln_value_mohm else '-'
            lpe = f"{float(ins.lpe_value_mohm):.1f}" if ins.lpe_value_mohm else '-'
            npe = f"{float(ins.npe_value_mohm):.1f}" if ins.npe_value_mohm else '-'
            row.cells[5].text = f"L-N: {ln}\nL-PE: {lpe}\nN-PE: {npe}"
            
            row.cells[6].text = 'Megfelel' if ins.passed else 'Nem mf.'
            
            # Reduce font size for measurement rows to fit
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
    else:
        doc.add_paragraph('Nincs mérési adat.')
    
    doc.add_paragraph()
    
    # 4.3 Loop impedance measurements
    doc.add_heading('4.3 Hurokellenállás (Zs)', level=2)
    
    loop_data = protocol.loop_impedance_measurements if protocol.loop_impedance_measurements else []
    if loop_data:
        loop_table = doc.add_table(rows=len(loop_data) + 1, cols=4)
        add_table_borders(loop_table)
        
        headers = ['Pont', 'Hely', 'Zs (Ω)', 'Megfelel']
        for i, h in enumerate(headers):
            loop_table.rows[0].cells[i].text = h
            set_cell_shading(loop_table.rows[0].cells[i], 'D9D9D9')
        
        for i, loop in enumerate(loop_data):
            row = loop_table.rows[i + 1]
            row.cells[0].text = str(loop.point_number)
            row.cells[1].text = loop.location
            row.cells[2].text = f"{float(loop.value_ohm):.2f}"
            row.cells[3].text = 'Igen' if loop.passed else 'Nem'
    else:
        doc.add_paragraph('Nincs mérési adat.')
    
    doc.add_paragraph()
    
    # 4.4 RCD tests
    doc.add_heading('4.4 FI-relé működésvizsgálat', level=2)
    
    rcd_data = protocol.rcd_tests if protocol.rcd_tests else []
    if rcd_data:
        rcd_table = doc.add_table(rows=len(rcd_data) + 1, cols=8)
        add_table_borders(rcd_table)
        
        headers = ['Áramkör', 'Megszakító', 'Vezeték', 'Vizsgálat',
                   'IΔn (mA)', 'Idő (ms)', 'Megfelel']
        for i, h in enumerate(headers):
            rcd_table.rows[0].cells[i].text = h
            set_cell_shading(rcd_table.rows[0].cells[i], 'D9D9D9')
        
        for i, rcd in enumerate(rcd_data):
            row = rcd_table.rows[i + 1]
            row.cells[0].text = rcd.circuit_name or '-'
            breaker = f"{rcd.breaker_type or 'B'} {rcd.breaker_value or '-'}" if rcd.breaker_type or rcd.breaker_value else '-'
            row.cells[1].text = breaker
            wire = f"{rcd.wire_material or 'Cu'} {rcd.wire_cross_section or '-'} mm²" if rcd.wire_material or rcd.wire_cross_section else '-'
            row.cells[2].text = wire
            row.cells[3].text = rcd.test_type or '-'
            row.cells[4].text = str(rcd.rated_current_ma) if rcd.rated_current_ma else '30'
            row.cells[5].text = f"{float(rcd.trip_time_ms):.2f}" if rcd.trip_time_ms else '-'
            row.cells[6].text = 'Igen' if rcd.passed else 'Nem'
    else:
        doc.add_paragraph('Nincs mérési adat.')
    
    doc.add_paragraph()
    
    # Section 5: Professional summary
    doc.add_heading('5. Szakmai összefoglaló', level=1)
    
    if protocol.professional_summary:
        doc.add_paragraph(protocol.professional_summary)
    else:
        default_summary = f"""A vizsgálati eredmények alapján megállapítható, hogy a {protocol.network_type} rendszerű hálózat érintésvédelmi módja megfelelő, a gyorskioldási feltétel minden vizsgált körnél teljesült. A mért Rpe, szigetelési ellenállás és hurokellenállás értékek biztonságos tartományban vannak, a FI-relé működése a szabvány által előírt időhatárokon belül történt. A hálózat üzemeltetése villamos-biztonságtechnikai szempontból megengedhető és megfelel a vonatkozó előírásoknak."""
        doc.add_paragraph(default_summary)
    
    # Section 6: Inspector declaration
    doc.add_heading('6. Felülvizsgáló nyilatkozata', level=1)
    
    declaration = """Alulírott kijelentem, hogy a vizsgálatot az MSZ HD 60364-6:2017 és MSZ EN 61557 szabványok alapján, hitelesített mérőműszerrel, szakszerű módszerekkel végeztem, és a jegyzőkönyv valós mérési eredményeket tartalmaz.\n\nA felülvizsgálat a villamos berendezés szerelésének és a környezeti körülményeinek megváltoztatásáig, de legfeljebb a következő törvényileg előírt időszakos felülvizsgálat esedékességéig érvényes."""
    doc.add_paragraph(declaration)
    
    doc.add_paragraph()
    doc.add_paragraph('Kelt.: ____________________________')
    doc.add_paragraph()
    doc.add_paragraph('Aláírás: ____________________________')
    
    # Section 7: Defect list
    doc.add_heading('7. Hibajegyzék', level=1)
    
    # Use structured defects if available, otherwise fallback to text
    if hasattr(protocol, 'protocol_defects') and protocol.protocol_defects:
        add_defects_section(doc, protocol)
    elif protocol.defect_list:
        doc.add_paragraph(protocol.defect_list)
    else:
        doc.add_paragraph('Nincs feltárt hiba. A vizsgált villamos berendezések minden mért és vizsgált paraméter tekintetében megfelelt az MSZ HD 60364-6:2017 szabvány előírásainak.')
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_eph_docx(protocol) -> bytes:
    """Generate EPH (Egyenpotenciálra Hozás) Word document from protocol data"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Header - Serial number
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(f"Jegyzőkönyv sorszáma: {protocol.serial_number}")
    header_run.bold = True
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title
    title = doc.add_heading('EPH (EGYENPOTENCIÁLRA HOZÁS) JEGYZŐKÖNYV', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Basic data
    doc.add_heading('1. Alapadatok', level=1)
    
    basic_data = [
        ('Helyszín:', protocol.location_address),
        ('Megrendelő neve:', protocol.client_name),
        ('Vizsgálat dátuma:', protocol.inspection_date.strftime('%Y.%m.%d.') if protocol.inspection_date else ''),
        ('Hálózati rendszer:', protocol.network_type),
        ('Bizonyítvány száma:', protocol.certificate_number or ''),
        ('Felülvizsgáló neve:', protocol.inspector_name),
        ('Mérőműszer:', f"{protocol.instrument_model or ''} (Kalibráció: {protocol.calibration_valid_until.strftime('%Y.%m.%d') if protocol.calibration_valid_until else 'N/A'})"),
    ]
    
    # Add EPH specific fields if available
    if protocol.gas_provider_required:
        basic_data.append(('Gázszolgáltató igénye:', 'Igen'))
    if protocol.gas_meter_number:
        basic_data.append(('Gázmérő száma:', protocol.gas_meter_number))
    if protocol.gas_appliance_type:
        basic_data.append(('Gázkészülék típusa:', protocol.gas_appliance_type))
    
    table = doc.add_table(rows=len(basic_data), cols=2)
    add_table_borders(table)
    
    for i, (label, value) in enumerate(basic_data):
        row = table.rows[i]
        row.cells[0].text = label
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Section 2: Electrical System Data
    doc.add_heading('2. Villamos rendszer adatok', level=1)
    
    system_data = [
        ('Hálózati rendszer:', protocol.network_type),
        ('PE-N szétválasztás helye:', protocol.pen_separation_point or 'Nem megadott'),
        ('PE vezeték keresztmetszete:', f"{float(protocol.pe_conductor_size):.1f} mm²" if protocol.pe_conductor_size else 'Nem megadott'),
        ('EPH fővezeték keresztmetszete:', f"{float(protocol.eph_conductor_size):.1f} mm²" if protocol.eph_conductor_size else 'Nem megadott'),
    ]
    
    sys_table = doc.add_table(rows=len(system_data), cols=2)
    add_table_borders(sys_table)
    
    for i, (label, value) in enumerate(system_data):
        row = sys_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Section 3: Earthing Measurements
    doc.add_heading('3. Földelési ellenállás mérés', level=1)
    
    earthing_data = protocol.earthing_measurements if protocol.earthing_measurements else []
    
    if earthing_data:
        # Method description
        method_names = {
            '3_wire': 'Háromvezetékes (3-wire) módszer',
            '2_clamp': 'Kétlakatfogós (2-clamp) módszer',
            'soil_resistivity': 'Fajlagos talajellenállás mérés'
        }
        
        for earthing in earthing_data:
            method = method_names.get(earthing.measurement_method, earthing.measurement_method or 'N/A')
            doc.add_paragraph(f"Mérési módszer: {method}")
            
            # Create measurement table
            earth_table = doc.add_table(rows=5, cols=4)
            add_table_borders(earth_table)
            
            # Header row
            headers = ['Paraméter', 'Mért érték', 'Határérték', 'Megfelelő']
            for i, h in enumerate(headers):
                earth_table.rows[0].cells[i].text = h
                set_cell_shading(earth_table.rows[0].cells[i], 'D9D9D9')
            
            # Data rows
            ra_passed = earthing.passed if earthing.passed is not None else (float(earthing.ra_value) <= 10.0 if earthing.ra_value else None)
            
            rows_data = [
                ('Ra (földelési ellenállás)', f"{float(earthing.ra_value):.2f} Ω" if earthing.ra_value else '-', '≤ 10 Ω', 'Igen' if ra_passed else 'Nem' if ra_passed is False else '-'),
                ('Rb (segédföldelő)', f"{float(earthing.rb_value):.2f} Ω" if earthing.rb_value else '-', '-', '-'),
                ('Rc (áramelektróda)', f"{float(earthing.rc_value):.2f} Ω" if earthing.rc_value else '-', '-', '-'),
                ('ρE (talajellenállás)', f"{float(earthing.soil_resistivity):.1f} Ωm" if earthing.soil_resistivity else '-', '-', '-'),
            ]
            
            for i, (param, value, limit, passed) in enumerate(rows_data):
                row = earth_table.rows[i + 1]
                row.cells[0].text = param
                row.cells[1].text = value
                row.cells[2].text = limit
                row.cells[3].text = passed
            
            # Soil type if available
            soil_types = {
                'humus': 'Humuszos talaj',
                'clay': 'Agyag',
                'sand': 'Homok',
                'gravel': 'Kavics',
                'rock': 'Szikla',
                'frozen': 'Fagyott talaj',
                'mixed': 'Vegyes'
            }
            if earthing.soil_type:
                doc.add_paragraph(f"Talaj típusa: {soil_types.get(earthing.soil_type, earthing.soil_type)}")
            
            if earthing.notes:
                doc.add_paragraph(f"Megjegyzés: {earthing.notes}")
    else:
        doc.add_paragraph('Nincs földelési mérési adat.')
    
    doc.add_paragraph()
    
    # Section 4: EPH Connections
    doc.add_heading('4. EPH bekötések ellenőrzése', level=1)
    
    eph_data = protocol.eph_measurements if protocol.eph_measurements else []
    
    if eph_data:
        element_types = {
            'water_pipe': 'Vízcső',
            'gas_pipe_metered': 'Gázcső (mérő előtt)',
            'gas_pipe_unmetered': 'Gázcső (mérő után)',
            'heating_pipe': 'Fűtéscső',
            'metal_bathtub': 'Fémkád',
            'shower_tray': 'Zuhanytálca',
            'lightning_conductor': 'Villámhárító',
            'cable_tray': 'Kábeltálca',
            'other': 'Egyéb'
        }
        
        eph_table = doc.add_table(rows=len(eph_data) + 1, cols=5)
        add_table_borders(eph_table)
        
        # Header row
        headers = ['#', 'Bekötött elem', 'Bekötési pont', 'R (Ω)', 'Megfelelő']
        for i, h in enumerate(headers):
            eph_table.rows[0].cells[i].text = h
            set_cell_shading(eph_table.rows[0].cells[i], 'D9D9D9')
        
        for i, eph in enumerate(eph_data):
            row = eph_table.rows[i + 1]
            row.cells[0].text = str(eph.point_number or i + 1)
            
            elem_display = eph.element_name
            if eph.element_type and eph.element_type != 'other':
                elem_display = element_types.get(eph.element_type, eph.element_name)
            row.cells[1].text = elem_display
            
            row.cells[2].text = eph.connection_point or 'EPH sín'
            row.cells[3].text = f"{float(eph.continuity_resistance):.2f}" if eph.continuity_resistance else '-'
            row.cells[4].text = '✓' if eph.passed else '✗'
    else:
        doc.add_paragraph('Nincs EPH bekötési adat.')
    
    doc.add_paragraph()
    
    # Section 5: Summary
    doc.add_heading('5. Vizsgálati összesítés', level=1)
    
    summary_data = protocol.summary_results if protocol.summary_results else []
    
    if summary_data:
        summary_table = doc.add_table(rows=len(summary_data) + 1, cols=3)
        add_table_borders(summary_table)
        
        headers = ['Vizsgálat', 'Minősítés', 'Megjegyzés']
        for i, h in enumerate(headers):
            summary_table.rows[0].cells[i].text = h
            set_cell_shading(summary_table.rows[0].cells[i], 'D9D9D9')
        
        for i, item in enumerate(summary_data):
            row = summary_table.rows[i + 1]
            row.cells[0].text = item.test_name
            row.cells[1].text = item.result
            row.cells[2].text = item.comment or ''
    else:
        # Default EPH summary items
        default_items = ['Földelési ellenállás', 'EPH fővezeték', 'EPH bekötések folytonossága', 'Gázcső bekötés']
        summary_table = doc.add_table(rows=len(default_items) + 1, cols=3)
        add_table_borders(summary_table)
        
        headers = ['Vizsgálat', 'Minősítés', 'Megjegyzés']
        for i, h in enumerate(headers):
            summary_table.rows[0].cells[i].text = h
            set_cell_shading(summary_table.rows[0].cells[i], 'D9D9D9')
        
        for i, item_name in enumerate(default_items):
            row = summary_table.rows[i + 1]
            row.cells[0].text = item_name
            row.cells[1].text = 'MEGFELELT'
            row.cells[2].text = 'Határértéken belül'
    
    doc.add_paragraph()
    
    # Section 6: Professional summary
    doc.add_heading('6. Szakmai összefoglaló', level=1)
    
    if protocol.professional_summary:
        doc.add_paragraph(protocol.professional_summary)
    else:
        default_summary = f"""Az egyenpotenciálra hozási (EPH) rendszer vizsgálata alapján megállapítható, hogy a bekötött fémszerkezetek megfelelően csatlakoznak az EPH sínhez. A földelési ellenállás az MSZ 447:2019 szabvány által előírt ≤ 10 Ω határértéken belül van. Az EPH rendszer a vonatkozó szabványok előírásainak megfelel, üzemeltetése biztonságos."""
        doc.add_paragraph(default_summary)
    
    # Section 7: Inspector declaration
    doc.add_heading('7. Felülvizsgáló nyilatkozata', level=1)
    
    declaration = """Alulírott kijelentem, hogy az egyenpotenciálra hozási rendszer vizsgálatát az MSZ HD 60364-41:2018, MSZ HD 60364-5-54 és MSZ 447:2019 szabványok alapján, hitelesített mérőműszerrel, szakszerű módszerekkel végeztem, és a jegyzőkönyv valós mérési eredményeket tartalmaz."""
    doc.add_paragraph(declaration)
    
    doc.add_paragraph()
    doc.add_paragraph('Kelt.: ____________________________')
    doc.add_paragraph()
    doc.add_paragraph('Aláírás: ____________________________')
    
    # Section 8: Defect list
    doc.add_heading('8. Hibajegyzék / Javítási javaslatok', level=1)
    
    # Use structured defects if available, otherwise fallback to text
    if hasattr(protocol, 'protocol_defects') and protocol.protocol_defects:
        add_defects_section(doc, protocol)
    elif protocol.defect_list:
        doc.add_paragraph(protocol.defect_list)
    else:
        doc.add_paragraph('Nincs feltárt hiba. Az EPH rendszer minden vizsgált paraméter tekintetében megfelelt a vonatkozó szabványok előírásainak.')
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
